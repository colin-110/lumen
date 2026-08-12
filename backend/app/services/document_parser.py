"""Document -> plain text extraction.

Uses PyMuPDF for PDF and python-docx for Word docs — both are pure/native
extensions with no external system binaries, which keeps container builds
fast and avoids OCR flakiness for the common "born-digital" document case.

Scanned/image-only PDF pages have no text layer for PyMuPDF to read, so
each page that yields near-nothing falls back to Tesseract OCR: rendered
to an image at OCR_DPI and passed through pytesseract. This is a per-page
fallback (not whole-document), so a PDF mixing real text pages with
scanned ones only pays the (much slower) OCR cost on the pages that
actually need it.
"""

from __future__ import annotations

import io
import logging
from pathlib import Path

import fitz  # PyMuPDF
from docx import Document as DocxDocument

from app.core.config import settings

logger = logging.getLogger(__name__)

_TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".log"}


class UnsupportedDocumentError(Exception):
    pass


def extract_text(file_path: str, filename: str) -> str:
    suffix = Path(filename).suffix.lower()

    if suffix == ".pdf":
        return _extract_pdf(file_path)
    if suffix == ".docx":
        return _extract_docx(file_path)
    if suffix in _TEXT_EXTENSIONS:
        return _extract_plain_text(file_path)

    # Last resort: try decoding as text before giving up, since some plain
    # text files arrive without a recognized extension. `_extract_plain_text`
    # no longer raises on undecodable bytes (it falls back through cp1252 and
    # latin-1, which accept anything), so the binary check has to be explicit
    # here — otherwise an unknown binary would be "successfully" ingested as
    # several thousand chunks of mojibake.
    if _looks_binary(file_path):
        raise UnsupportedDocumentError(f"Unsupported file type: {suffix or '(none)'}")
    return _extract_plain_text(file_path)


def _looks_binary(file_path: str, sample_bytes: int = 8192) -> bool:
    """A NUL byte in the first few KB is the standard heuristic, and the one
    `file(1)` and git both use. Text formats do not contain them."""
    with open(file_path, "rb") as f:
        return b"\x00" in f.read(sample_bytes)


def _extract_pdf(file_path: str) -> str:
    parts = []
    ocr_pages = 0
    with fitz.open(file_path) as doc:
        for page in doc:
            text = page.get_text("text")
            if settings.OCR_ENABLED and len(text.strip()) < settings.OCR_MIN_CHARS_PER_PAGE:
                ocr_text = _ocr_page(page)
                if len(ocr_text.strip()) > len(text.strip()):
                    text = ocr_text
                    ocr_pages += 1
            parts.append(text)
    if ocr_pages:
        logger.info("OCR fallback used on %d page(s) of %s", ocr_pages, file_path)
    return "\n\n".join(parts)


def _ocr_page(page: fitz.Page) -> str:
    """Render a page to an image and run Tesseract on it. Imported lazily so
    a missing/broken tesseract install only breaks OCR, not every upload —
    born-digital PDFs (the common case) never touch this code path."""
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        logger.warning("pytesseract/Pillow not installed; skipping OCR fallback")
        return ""

    try:
        pix = page.get_pixmap(dpi=settings.OCR_DPI)
        image = Image.open(io.BytesIO(pix.tobytes("png")))
        return pytesseract.image_to_string(image)
    except Exception:
        logger.warning("OCR fallback failed for a page; leaving it blank", exc_info=True)
        return ""


def _extract_docx(file_path: str) -> str:
    doc = DocxDocument(file_path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        parts.append(_render_table(table))
    return "\n\n".join(p for p in parts if p.strip())


def _render_table(table) -> str:
    """Render a table as markdown, repeating the header on every row group.

    Flattening every row to "a | b | c" dropped which column each value was
    in, so a chunk taken from the middle of a table reached the model as a
    row of bare values with no idea what they meant. Keeping the header row
    and marking it as such lets the model attribute a cell to its column, and
    the pipe layout survives chunking as plain text.
    """
    rows = [[cell.text.strip().replace("\n", " ") for cell in row.cells] for row in table.rows]
    rows = [r for r in rows if any(c for c in r)]
    if not rows:
        return ""
    if len(rows) == 1:
        return " | ".join(rows[0])

    header, *body = rows
    lines = [
        " | ".join(header),
        " | ".join("---" for _ in header),
        *(" | ".join(r) for r in body),
    ]
    return "\n".join(lines)


def _extract_plain_text(file_path: str) -> str:
    """Decode a text file, preferring UTF-8 but never failing on it.

    Opening with a strict UTF-8 codec meant a CSV exported from Excel
    (cp1252) or any Latin-1 document raised UnicodeDecodeError and was
    recorded as unprocessable — for the sake of a handful of accented
    characters. Try the likely encodings in order, then fall back to lossy
    UTF-8: a document with a few replacement characters is still searchable,
    a rejected one is not.
    """
    raw = Path(file_path).read_bytes()
    if raw.startswith(b"\xef\xbb\xbf"):
        return raw.decode("utf-8-sig")

    for encoding in ("utf-8", "cp1252", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")
